#!/usr/bin/env python3
"""Assemble the translated-book Markdown into a SimSun-formatted Word document."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn


def set_font(run, name="宋体", size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)


def first_line_indent(p, chars=2):
    """Set first-line indent to N full-width character widths at 12pt."""
    pf = p.paragraph_format
    pf.first_line_indent = Pt(12 * chars)


def main():
    md_path = Path(sys.argv[1])
    out_docx = Path(sys.argv[2])
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    # ---- Front matter ----
    book_title = ""
    subtitle = ""
    author_line = ""
    note_line = ""
    toc_titles = []

    # Locate front-matter pieces and TOC section titles.
    lines_clean = [ln.rstrip() for ln in lines]
    i = 0
    n = len(lines_clean)
    body_start = 0
    while i < n:
        ln = lines_clean[i]
        if ln.startswith('# ') and not book_title:
            book_title = ln[2:].strip()
        elif ln.startswith('## '):
            subtitle = ln[3:].strip()
        elif ln.startswith('**作者：') or ln.startswith('**作者:**'):
            author_line = ln.replace('**', '').strip()
        elif ln.startswith('**说明：') or ln.startswith('**说明:**'):
            note_line = ln.replace('**', '').strip()
        elif ln.strip() == '# 目录':
            # collect TOC entries until first '---' after heading
            j = i + 1
            while j < n and not lines_clean[j].startswith('# '):
                t = lines_clean[j].strip()
                if re.match(r'^\d+\.\s', t):
                    toc_titles.append(re.sub(r'^\d+\.\s*', '', t))
                j += 1
            body_start = j
            break
        i += 1

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(book_title), size=18, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(subtitle), size=14, bold=True)
    p.paragraph_format.space_after = Pt(12)

    if author_line:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(author_line), size=12)
        p.paragraph_format.space_after = Pt(12)

    if note_line:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(note_line), size=10.5, color=RGBColor(0x88, 0x88, 0x88))
        p.paragraph_format.space_after = Pt(18)

    # ---- TOC ----
    if toc_titles:
        p = doc.add_paragraph()
        set_font(p.add_run("目录"), size=16, bold=True)
        p.paragraph_format.space_after = Pt(6)
        for idx, t in enumerate(toc_titles, 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(24)
            p.paragraph_format.space_after = Pt(2)
            set_font(p.add_run(f"{idx}. {t}"), size=12)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---- Body ----
    # Parse sections after the TOC marker.
    cur_section = None
    for idx in range(body_start, n):
        ln = lines_clean[idx]
        if not ln.strip():
            continue
        if ln.startswith('# '):
            cur_section = ln[2:].strip()
            p = doc.add_paragraph()
            set_font(p.add_run(cur_section), size=16, bold=True)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif ln.startswith('---'):
            continue
        else:
            # body paragraph
            p = doc.add_paragraph()
            first_line_indent(p)
            set_font(p.add_run(ln), size=12)

    # ---- Footer disclaimer ----
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_font(p.add_run("本文档为 AI 辅助个人学习译稿，非正式出版物，仅供个人学习使用；学术引用请以英文原版为准。"),
             size=10.5, color=RGBColor(0x88, 0x88, 0x88))

    doc.save(out_docx)
    print("WROTE", out_docx)


if __name__ == "__main__":
    main()
