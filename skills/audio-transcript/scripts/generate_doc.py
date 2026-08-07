#!/usr/bin/env python3
"""Generate a formatted Word transcript with SimSun font."""
import json, sys, datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, name="宋体", size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)

def fmt_ts(sec):
    sec = int(round(sec))
    return f"{sec//3600:02d}:{sec%3600//60:02d}:{sec%60:02d}"

def main():
    turns_json = sys.argv[1]
    out_docx = sys.argv[2]
    title = sys.argv[3] if len(sys.argv) > 3 else "对话稿"
    subtitle = sys.argv[4] if len(sys.argv) > 4 else ""
    footer = sys.argv[5] if len(sys.argv) > 5 else ""

    with open(turns_json, encoding="utf-8") as f:
        turns = json.load(f)

    doc = Document()

    # Default style -> SimSun, 12pt, 1.5 line spacing
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(title), size=16, bold=True)
    p.paragraph_format.space_after = Pt(4)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(subtitle), size=10.5, color=RGBColor(0x88, 0x88, 0x88))
    p.paragraph_format.space_after = Pt(12)

    for i, t in enumerate(turns, 1):
        # Speaker + timestamp line
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(f"说话人{t['speaker']}"), size=12, bold=True)
        set_font(p.add_run(f"  [{fmt_ts(t['start'])}]"), size=10.5, color=RGBColor(0x66, 0x66, 0x66))

        # Content line
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Pt(24)
        p2.paragraph_format.line_spacing = 1.5
        p2.paragraph_format.space_after = Pt(6)
        set_font(p2.add_run(t['text']), size=12)

    if footer:
        doc.add_paragraph()
        p = doc.add_paragraph()
        set_font(p.add_run(footer), size=10.5, color=RGBColor(0x88, 0x88, 0x88))

    doc.save(out_docx)
    print("WROTE", out_docx)

if __name__ == "__main__":
    main()
